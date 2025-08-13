import streamlit as st
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from PIL import Image
from dotenv import load_dotenv
import fitz  # PyMuPDF
import os
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

# 🔐 Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)

# ⚡ Gemini multimodal model
model = genai.GenerativeModel("gemini-1.5-flash")

# 🧠 Initialize session state for chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# 📄 Sidebar for file upload
st.sidebar.title("Upload Files")
uploaded_files = st.sidebar.file_uploader(
    "Upload text, PDF, image, or Word files",
    type=["txt", "pdf", "jpg", "png", "docx"],
    accept_multiple_files=True
)

# 🌐 Website input
url_input = st.sidebar.text_input("Enter a website URL")

# 📚 Document store setup
documents = []
images = []

# 🔧 Chunking setup
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

# 📥 Process uploaded files
if uploaded_files:
    for uploaded_file in uploaded_files:
        file_type = uploaded_file.type
        text = ""

        if file_type == "text/plain":
            text = uploaded_file.read().decode("utf-8")
        elif file_type == "application/pdf":
            pdf_reader = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            for page in pdf_reader:
                text += page.get_text()
        elif file_type.startswith("image"):
            image = Image.open(uploaded_file)
            images.append(image)
            st.image(image, caption="Uploaded Image", use_column_width=True)
            continue
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            docx_file = DocxDocument(uploaded_file)
            text = "\n".join([para.text for para in docx_file.paragraphs])

        if text:
            chunks = splitter.split_text(text)
            documents.extend([Document(page_content=chunk) for chunk in chunks])

# 🌐 Process website URL
if url_input:
    try:
        response = requests.get(url_input)
        soup = BeautifulSoup(response.text, 'html.parser')
        web_text = ' '.join([p.get_text() for p in soup.find_all('p')])
        chunks = splitter.split_text(web_text)
        documents.extend([Document(page_content=chunk) for chunk in chunks])
    except Exception as e:
        st.error(f"Failed to fetch website: {e}")

# 🧠 Create vector store
if documents:
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    db = FAISS.from_documents(documents, embeddings)

# 💬 User query
query = st.text_input("Ask a question about your uploaded content")

if query:
    # 🔍 RAG context from vector store
    context = ""
    if documents:
        docs = db.similarity_search(query, k=3)
        context = "\n".join([doc.page_content for doc in docs])

    # 🧠 Include chat history in context
    history_context = "\n".join(
        [f"User: {turn['user']}\nBot: {turn['bot']}" for turn in st.session_state.chat_history]
    )

    full_context = f"{history_context}\n{context}" if history_context else context

    # 🤖 Multimodal input: query + context + images
    inputs = [query]
    if full_context:
        inputs.append(full_context)
    inputs.extend(images)

    response = model.generate_content(inputs)

    # 💾 Save to chat history
    st.session_state.chat_history.append({"user": query, "bot": response.text})

# 🗂️ Display chat history
if st.session_state.chat_history:
    st.markdown("### 🗂️ Chat History")
    for turn in st.session_state.chat_history:
        st.markdown(f"**You:** {turn['user']}")
        st.markdown(f"**Gemini:** {turn['bot']}")

# 🧹 Clear history button
if st.button("Clear Chat History"):
    st.session_state.chat_history = []
